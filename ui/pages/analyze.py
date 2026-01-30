import time
import docx
import docx2txt
from dataclasses import dataclass

file = "C:\\Users\\cyrille.faucon\\.my_app\\CV_n_1819877.docx"

@dataclass
class Experience :
    title    : str
    duration : str
    skills   : list


def analyze_cv(cv) :
    doc = docx.Document(file)
    
    read_exp      = False
    read_title    = False
    read_company  = False
    read_duration = False 
    read_skills   = False

    title    = ""
    company  = ""
    duration = ""
    skills   = []
    loop_para = 0
    while loop_para < len(doc.paragraphs) :
        # print(loop_line)
        para = doc.paragraphs[loop_para]

        if "Heading" in para.style.name :
            if "Expériences" in para.text :
                read_exp = True

        elif read_exp == True and read_title == False and line != "" :
            read_title = True
            title = line
            loop_line += 1

        elif read_exp == True and read_title == True and read_company == False and line != "" :
            read_company = True
            company = line
            loop_line += 1

        elif read_exp == True and read_title == True and read_company == True and read_duration == False and line != "" :
            read_duration = True
            duration = line
            loop_line += 1

        elif read_exp == True and read_title == True and read_company == True and read_duration == True and read_skills == False and line != "" and "Environnement technique" in line :
            read_skills = True
            skills_line = loop_line + 1
            while skills_line < len(lines) and "+" in lines[skills_line] :
                skills.append(lines[skills_line])
                skills_line += 1
            if len(skills) == 0 :
                loop_line += 1
            else :
                loop_line = skills_line
               
        else :
            loop_line += 1
        
        if read_exp == True and read_title == True and read_company == True and read_duration == True and read_skills == True :
            print(title)
            print(company)
            print(duration)
            print(skills)
            skills = []
            read_title    = False
            read_company  = False
            read_duration = False 
            read_skills   = False

            
            
    return True

# analyze_cv(file)


doc = docx.Document(file)
# print(doc.paragraphs)

for p in doc.paragraphs :
    if "Heading" in p.style.name : 
        print(p.style.name + "   " + p.text)
