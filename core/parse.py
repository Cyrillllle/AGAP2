import time
import docx
import re
import sqlite3
from dataclasses import dataclass


@dataclass
class Experience :
    title     : str
    company   : str
    duration  : str
    details   : list



Heading_dict_type_1 = {
    "Heading 2" : "title",
    "Heading 6" : "company",
    "Heading 7" : "duration",
    "details"   : "Heading 5"
}

Heading_dict_type_2 = {
    "Heading 2" : "title",
    "Heading 5" : "company",
    "Heading 6" : "duration", 
    "details"   : "Heading 4"
}

def get_details(paragraphs, index, heading_dict, exp_details : Experience) :
    while index < len(paragraphs) :
        para = paragraphs[index]
        para_style = para.style.name
        if heading_dict == Heading_dict_type_1 and "Heading 4" in para_style  :
            index += 1
        elif "Heading" in para_style : 
            break
        else : 
            if para.text != "" and not re.search("\\d / \\d", para.text) :
                exp_details.details[-1].append(para.text)
            index += 1
    return index


def parse_exp_details(paragraphs, index, heading_dict, exp_details : Experience) :
    while index < len(paragraphs) :
        para = paragraphs[index]
        para_style = para.style.name
        if para_style in heading_dict :
            if heading_dict[para_style] == "title" :
                break
            elif heading_dict[para_style] == "company" :
                exp_details.company = para.text
            elif heading_dict[para_style] == "duration" :
                exp_details.duration = para.text
            index += 1
        elif heading_dict["details"] in para_style :
            exp_details.details.append([para.text])
            new_index = get_details(paragraphs, index + 1, heading_dict, exp_details)
            index = new_index
        else :
            index += 1


def parse_skills(paragraphs, index) :
    skills = []
    while index < len(paragraphs) :
        para = paragraphs[index]
        para_style = para.style.name
        if "Heading 1" in para_style :
            break
        else :
            if "+" in para.text :
                skills.append(para.text.strip("+ \t"))
        index += 1
    return skills

                
def parse_cv(file) :
    doc = docx.Document(file)
    reading_exp = -1
    reading_skills = -1
    heading_dict = Heading_dict_type_2
    experiences = []
    skills = []
    for index, p in enumerate(doc.paragraphs) :
        if "Heading 4" in p.style.name and p.text == "THINK2MORROW" :
            heading_dict = Heading_dict_type_1
            break
    for index, p in enumerate(doc.paragraphs) :
        if reading_exp == 0 and p.style.name in heading_dict and heading_dict[p.style.name] == "title":
            exp_details = Experience(p.text, "", "", [])
            parse_exp_details(doc.paragraphs, index + 1, heading_dict, exp_details)
            experiences.append(exp_details)

        if "Heading 1" in p.style.name and "Expériences" in p.text : 
            reading_exp = 0
            continue

        if "Heading 1" in p.style.name and "Compétences" in p.text : 
            skills = parse_skills(doc.paragraphs, index + 1)
            print(skills)
            continue
                
        if reading_exp == 0 and "Heading 1" in p.style.name :
            reading_exp = 1

        if reading_skills == 0 and "Heading 1" in p.style.name :
            reading_skills = 1

        if reading_exp == 1 and reading_skills == 1 :
            break
    return experiences, skills
