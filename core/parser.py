import streamlit as st
import pathlib
import os
import time
import docx
import re
import sqlite3
from queue import Queue
import json
import shelve
import time
import threading
import io
import json
from unidecode import unidecode
from datetime import datetime
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from streamlit_autorefresh import st_autorefresh

from api.client import api_request, RequestType, GetAllUsers, GetUserCv, ExportCv
from core.storage import USER_DATA_DIR, TOKEN_PATH
from core.database import *
from core.parser import *


@dataclass
class Experience :
    title     : str
    company   : str
    duration  : str
    details   : list


def get_details_heading_type(paragraphs) :
    heading_type = ""
    for p in paragraphs :
        if ("taches effectuees" == unidecode(p.text.lower()) or "contexte" == unidecode(p.text.lower()) or "tasks completed" == unidecode(p.text.lower()) or "technical environment" == unidecode(p.text.lower())) :
            heading_type = p.style.name
            break
    return heading_type

def get_title_heading_type(paragraphs) :
    heading_type = ""
    for p in paragraphs :
        if "Heading" in p.style.name and ("experience" in unidecode(p.text.lower()) or "competence" in unidecode(p.text.lower())) and ("an" not in unidecode(p.text.lower()) and "dossier" not in unidecode(p.text.lower())) :
            heading_type = p.style.name
            break
    return heading_type

def get_exp_details(paragraphs, index, details_heading, company_heading, title_heading, exp_details : Experience) :
    sub_text = []
    while index < len(paragraphs) :
        para = paragraphs[index]
        para_style = para.style.name
        if (company_heading in para_style or title_heading in para_style) and para.text != "THINK2MORROW" and not re.search("\\d / \\d", para.text) and para.text != "" : 
            exp_details.details.append(sub_text)
            sub_text = []
            break
        elif details_heading in para_style and para.text != "THINK2MORROW" :
            if len(sub_text) != 0 :
                exp_details.details.append(sub_text)
                sub_text = [para.text]
            elif para.text != "" and para.text != "THINK2MORROW" :
                sub_text.append(para.text)
        else : 
            if para.text != "" and not re.search("\\d / \\d", para.text) and para.text != "THINK2MORROW" :
                # if len(sub_text) != 0 and "techni" in unidecode(sub_text[0].lower()) :
                text = para.text.strip("+ \t")
                sub_text += (text.split(","))
        index += 1
        # print(sub_text)
    if len(sub_text) != 0 :
        exp_details.details.append(sub_text)
    return index

def parse_skills(paragraphs, index, title_heading) :
    skills = []
    while index < len(paragraphs) :
        para = paragraphs[index]
        para_style = para.style.name
        if title_heading in para_style :
            break
        else :
            if "Heading" not in para.style.name and para.text != "" and para.text != "THINK2MORROW" and not re.search("\\d / \\d", para.text) :
                skills.append(para.text.strip("+ \t"))
        index += 1
    return skills, index

def get_company_heading(paragraphs, index) :
    while index < len(paragraphs) :
        p = paragraphs[index]
        if p.text != "" :
            break
        else : 
            index += 1

    return p.style.name, index

def parse_cv(file, cv_id) :
    doc = docx.Document(io.BytesIO(file))
    # doc = docx.Document(file)
    reading_exp = -1
    reading_skills = -1
    title_heading  = get_title_heading_type(doc.paragraphs)
    details_heading = get_details_heading_type(doc.paragraphs)
    company_heading = ""
    index = 0
    experiences = []
    skills = []
    while index < len(doc.paragraphs) :
        p = doc.paragraphs[index]
        if title_heading in p.style.name and "experience" in unidecode(p.text.lower()) and "resume" not in unidecode(p.text.lower()) and "summary" not in unidecode(p.text.lower()) : 
            reading_exp = 0
            company_heading, oIndex = get_company_heading(doc.paragraphs, index+1)
            index = oIndex - 1
        elif reading_exp == 0 :
            if company_heading in p.style.name and "THINK2MORROW" not in p.text and not re.search("\\d / \\d", p.text) and p.text != "" :
                exp_details = Experience(doc.paragraphs[index].text, doc.paragraphs[index+1].text, doc.paragraphs[index+2].text, [])
                oIndex = get_exp_details(doc.paragraphs, index+3, details_heading, company_heading, title_heading, exp_details)
                index = oIndex - 1
                # exp.append(exp_details)
                exp_dict = asdict(exp_details)
                experiences.append(exp_dict)
            elif title_heading in p.style.name :
                reading_exp = 1
        if title_heading in p.style.name and ("competences" == unidecode(p.text.lower()) or "skills" == unidecode(p.text.lower())) : 
            skills, oIndex = parse_skills(doc.paragraphs, index+1, title_heading)
            index = oIndex - 1
        index += 1
    exp_json = json.dumps(experiences)
    skills_json = json.dumps(skills)
    # print(experiences)
    return exp_json, skills_json


def parse_worker(pipelineManager, selection, writer_queue):
    treated = 0
    total = get_total_cv_parsing()
    total_time = 0

    for cv_id, cv_raw in read_raw_data():
        try :
            # if pipelineManager._stop_event :
            #     break
            start_time = time.time()
            exp, skills = parse_cv(cv_raw, cv_id)
            elapsed_time = time.time() - start_time
            treated += 1
            total_time = total_time + elapsed_time
            mean_treatmment_time = total_time / treated
            print(mean_treatmment_time)
            remaining_time = mean_treatmment_time * (total - treated)
            if remaining_time >= 60 :
                estimation_mn = str(int(remaining_time/60))+"mn"
            else : 
                estimation_mn = ""
            estimation_sec = str(int((remaining_time%60))) + "s"
            estimation = estimation_mn + estimation_sec
            
            pipelineManager.progress = treated / total      
            pipelineManager.message = f"{treated}/{total} profils traités. Environ {estimation} restantes"
            writer_queue.put({"type": "upsert_cv_parsed", "data": (cv_id, exp, skills)})
            mark_cv_parsed(cv_id)
        except Exception as e :
            print(e)

    
    pipelineManager.step += 1


def start_parse(pipelineManager, stop_event):
    init_db()
    existing_users = load_users_cv_dates()
    writer_queue, stop_event, writer_thread = start_writer()
    parse_worker(pipelineManager, stop_event, writer_queue)


# path = "C:\\Users\\cyrille.faucon\\.my_app\\data.db-x-cv_raw-4561980-data_raw.docx"

# parse_cv(path, 10)